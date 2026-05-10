import io
import json
import re
from difflib import SequenceMatcher

import google.generativeai as genai
import requests
import streamlit as st
from bs4 import BeautifulSoup
from docx import Document

# ─── Config ───────────────────────────────────────────────────────────────────

MODEL = "gemini-2.5-flash"

_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"
    ),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.9",
}

_LINKEDIN_SELECTORS = [
    "div.description__text",
    "div.show-more-less-html__markup",
    "div[class*='description']",
    "section.description",
]

_NAUKRI_SELECTORS = [
    "div.job-desc",
    "div.dang-inner-html",
    "section.job-desc",
    "div[class*='jd-desc']",
    "div[class*='jobDescription']",
]

_RESET_KEYS = {"analyzed", "result", "cv_text", "jd_text", "cv_generated", "word_bytes"}

_RE_STRIP_FENCES = re.compile(r'^```(?:json)?\s*|\s*```\s*$', re.MULTILINE)
# Used only for bullet-prefix detection during paragraph patching
_RE_BULLET_SYM = re.compile(r'^[•\-\*▪◦✓]\s')
_RE_BULLET_NUM = re.compile(r'^\d+[.)]\s')

# ─── DOCX helpers ─────────────────────────────────────────────────────────────

def _iter_paragraphs(doc: Document):
    """Yield every paragraph in the document, including those inside tables."""
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


def _strip_bullet(text: str) -> str:
    """Strip a leading bullet character for fuzzy-comparison purposes."""
    s = _RE_BULLET_SYM.sub("", text.strip())
    return _RE_BULLET_NUM.sub("", s).strip()


def _patch_paragraph(para, dst: str) -> None:
    """
    Replace the text content of a paragraph while leaving every run's
    character formatting (font, size, bold, italic, colour, hyperlink) intact.

    Strategy:
    - Strip any bullet char Gemini may have included in dst.
    - If the paragraph has a manual bullet prefix in its first run, keep that
      run verbatim and put the new content in the second run (clearing the rest).
    - If the bullet is in one single run, rebuild that run as prefix + new content.
    - If no bullet prefix (Word auto-list or plain paragraph), replace the first
      run's text and blank-out all others.
    """
    runs = para.runs
    if not runs:
        para.add_run(dst)
        return

    orig = para.text
    dst_content = _strip_bullet(dst) or dst.strip()

    m_sym = _RE_BULLET_SYM.match(orig)
    m_num = _RE_BULLET_NUM.match(orig)

    if (m_sym or m_num) and len(runs) >= 2:
        # Bullet char lives in its own first run (e.g. "• " then the text)
        runs[1].text = dst_content
        for run in runs[2:]:
            run.text = ""

    elif (m_sym or m_num) and len(runs) == 1:
        # Bullet char and text are in the same single run
        prefix = (m_sym or m_num).group(0)
        runs[0].text = prefix + dst_content

    else:
        # Word auto-list paragraph (bullet not in raw text) or plain line
        runs[0].text = dst_content
        for run in runs[1:]:
            run.text = ""

# ─── Core pipeline ────────────────────────────────────────────────────────────

def fetch_job_description(url: str) -> tuple[str | None, str | None]:
    try:
        resp = requests.get(url, headers=_HEADERS, timeout=15)
        resp.raise_for_status()
    except requests.exceptions.Timeout:
        return None, "Request timed out. The site may be slow or blocking scrapers."
    except requests.exceptions.HTTPError as e:
        return None, f"HTTP {e.response.status_code}: {e}"
    except requests.RequestException as e:
        return None, f"Network error: {e}"

    soup = BeautifulSoup(resp.content, "html.parser")
    for tag in soup(["script", "style", "nav", "header", "footer", "aside"]):
        tag.decompose()

    selectors = _LINKEDIN_SELECTORS if "linkedin.com" in url else (
        _NAUKRI_SELECTORS if "naukri.com" in url else []
    )
    for selector in selectors:
        el = soup.select_one(selector)
        if el:
            text = el.get_text(separator="\n", strip=True)
            if len(text) > 100:
                return text[:7000], None

    best_text = max(
        (d.get_text(separator="\n", strip=True) for d in soup.find_all("div")),
        key=len, default="",
    )
    if len(best_text) > 100:
        return best_text[:7000], None

    return None, (
        "Could not extract the job description. LinkedIn may require login. "
        "Try pasting the JD text manually."
    )


def extract_docx_text(docx_file) -> str:
    """Extract plain text from a DOCX, one line per paragraph."""
    doc = Document(docx_file)
    return "\n".join(para.text for para in _iter_paragraphs(doc))


def analyze_and_rewrite(cv_text: str, jd_text: str, api_key: str) -> dict:
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(
        model_name=MODEL,
        system_instruction=(
            "You are an expert CV writer and career coach. "
            "Respond with valid JSON only — no markdown fences, no prose."
        ),
        generation_config=genai.GenerationConfig(temperature=0),
    )
    prompt = f"""Compare this CV against the job description. Your job:

1. Extract every bullet point from the CV (lines beginning with •, -, *, ▪, ◦, ✓,
   a digit+period/paren, or any other list marker). Keep the exact original text.
2. For each bullet decide: is it **weak** (vague, irrelevant, lacks metrics, misaligned
   with the JD) or **strong** (specific, relevant, ideally quantified)?
   Be consistent and thorough — evaluate every single bullet point in the CV without
   skipping any. A bullet is weak if it is vague, lacks metrics, or doesn't use keywords
   from the JD. Do not skip bullets just because they seem acceptable.
3. Rewrite ONLY the weak bullets — improve phrasing, strengthen impact language, and
   naturally weave in missing JD keywords where genuinely applicable to the candidate's
   existing experience. Do NOT invent skills or experiences the CV doesn't mention.
   If the original starts with an action verb, the rewrite must also start with an action verb.
4. For each weak bullet, assign an impact_score from 1–10: how much does accepting this
   rewrite improve the JD match? 10 = addresses a critical required skill/keyword absent
   from the CV; 1 = minor phrasing polish with little strategic effect.
5. Provide a holistic JD match score (0–100) before and after ALL rewrites are applied.
6. Compute an ATS keyword coverage score (0–100):
   - Identify the 20–30 most important keywords/phrases from the JD.
   - ats_score_before = (keywords found in the original CV / total JD keywords) × 100, rounded.
   - ats_score_after  = (keywords found after ALL rewrites applied / total JD keywords) × 100, rounded.
7. Extract the job title and hiring company name from the JD.
   - job_title: the role being advertised (e.g. "Product Manager"). Required.
   - company_name: the hiring organisation (e.g. "Google"). Use null if not mentioned.

Return ONLY this JSON structure — no extra keys, no comments:
{{
  "job_title": "<role title>",
  "company_name": "<company name or null>",
  "match_score_before": <integer>,
  "match_score_after": <integer>,
  "ats_score_before": <integer 0–100>,
  "ats_score_after": <integer 0–100>,
  "bullets": [
    {{
      "original": "<exact bullet text>",
      "rewritten": "<improved bullet, or same as original if strong>",
      "is_weak": <true|false>,
      "reason": "<one sentence>",
      "impact_score": <integer 1–10 when is_weak is true, omit when is_weak is false>
    }}
  ]
}}

━━━ JOB DESCRIPTION ━━━
{jd_text}

━━━ CANDIDATE CV ━━━
{cv_text}"""

    raw = _RE_STRIP_FENCES.sub("", model.generate_content(prompt).text.strip())
    return json.loads(raw)


def apply_changes_to_docx(docx_file, accepted: list[dict]) -> bytes:
    """
    Open the original DOCX and surgically patch only the accepted bullet paragraphs.
    Every other paragraph — its text, formatting, spacing, hyperlinks — is untouched.
    """
    doc = Document(docx_file)
    all_paras = list(_iter_paragraphs(doc))

    for change in accepted:
        src = change["original"].strip()
        dst = change["rewritten"].strip()
        src_clean = _strip_bullet(src)

        best_para, best_ratio = None, 0.0

        for para in all_paras:
            pt = para.text.strip()
            if not pt:
                continue
            # Exact match (with or without bullet prefix)
            if pt == src or _strip_bullet(pt) == src_clean:
                best_para, best_ratio = para, 1.0
                break
            ratio = SequenceMatcher(None, src_clean, _strip_bullet(pt)).ratio()
            if ratio > best_ratio:
                best_ratio, best_para = ratio, para

        if best_para is not None and best_ratio >= 0.75:
            _patch_paragraph(best_para, dst)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# ─── Page setup ───────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="Bespoke CV",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
.bullet-original {
    background: #fff0f0;
    border-left: 4px solid #ef5350;
    padding: 10px 14px;
    border-radius: 6px;
    color: #b71c1c;
    text-decoration: line-through;
    font-size: 0.93rem;
    line-height: 1.55;
}
.bullet-kept {
    background: #fafafa;
    border-left: 4px solid #bdbdbd;
    padding: 10px 14px;
    border-radius: 6px;
    color: #616161;
    font-size: 0.93rem;
    line-height: 1.55;
}
.bullet-rewritten {
    background: #f0fff4;
    border-left: 4px solid #43a047;
    padding: 10px 14px;
    border-radius: 6px;
    color: #1b5e20;
    font-size: 0.93rem;
    line-height: 1.55;
}
.bullet-reason {
    font-size: 0.76rem;
    color: #9e9e9e;
    margin-top: 5px;
    font-style: italic;
}
.col-header {
    font-weight: 700;
    font-size: 0.85rem;
    text-transform: uppercase;
    letter-spacing: 0.06em;
    padding-bottom: 4px;
    border-bottom: 2px solid #e0e0e0;
    margin-bottom: 12px;
}
.impact-badge {
    display: inline-block;
    background: #fff3e0;
    color: #e65100;
    border-radius: 4px;
    padding: 1px 6px;
    font-size: 0.72rem;
    font-weight: 700;
    margin-right: 6px;
    vertical-align: middle;
}
</style>
""", unsafe_allow_html=True)

# ─── Sidebar ──────────────────────────────────────────────────────────────────

with st.sidebar:
    st.header("Configuration")
    gemini_api_key = st.text_input(
        "Gemini API Key",
        type="password",
        placeholder="AIza…",
        help="Used only for this session, never stored.",
    )
    st.caption(
        "Get your **free** key at "
        "[aistudio.google.com](https://aistudio.google.com/app/apikey)"
    )
    st.divider()
    st.caption(f"Model: `{MODEL}`")

# ─── Header ───────────────────────────────────────────────────────────────────

st.title("📄 Bespoke CV")
st.caption("Upload your CV as a Word file and a job posting URL — Gemini rewrites weak bullets to match the role.")
st.divider()

# ─── Inputs ───────────────────────────────────────────────────────────────────

left, right = st.columns(2)
with left:
    job_url = st.text_input(
        "Job Posting URL",
        placeholder="https://www.linkedin.com/jobs/view/…  or  https://www.naukri.com/…",
    )
with right:
    uploaded_cv = st.file_uploader(
        "Your CV (Word .docx)",
        type=["docx"],
        help="Upload your CV as a Word document (.docx). This allows the output to exactly match your original formatting.",
    )

with st.expander("Or paste the job description manually (overrides URL fetch)"):
    manual_jd = st.text_area("Job description text", height=160,
        placeholder="Paste the full JD here if the URL cannot be scraped…")

st.write("")
go = st.button("✦ Analyze & Optimize", type="primary", use_container_width=True)

# ─── Analysis pipeline ────────────────────────────────────────────────────────

if go:
    if not gemini_api_key.strip():
        st.error("Please enter your Gemini API key in the sidebar.")
        st.stop()
    if not job_url and not manual_jd.strip():
        st.error("Please enter a job posting URL or paste the job description.")
        st.stop()
    if not uploaded_cv:
        st.error("Please upload your CV as a .docx Word file.")
        st.stop()

    for k in list(st.session_state.keys()):
        if k.startswith("accept_") or k in _RESET_KEYS:
            del st.session_state[k]

    if manual_jd.strip():
        jd_text = manual_jd.strip()
        st.toast("Using manually pasted job description.", icon="📋")
    else:
        with st.status("Fetching job description from URL…", expanded=False) as status:
            jd_text, err = fetch_job_description(job_url)
            if err or not jd_text:
                status.update(label="Failed to fetch job description", state="error")
                st.error(err or "No content extracted from URL.")
                st.info("Tip: LinkedIn often blocks scrapers — paste the JD manually above.")
                st.stop()
            status.update(label=f"Job description fetched ({len(jd_text):,} chars)", state="complete")

    with st.status("Reading your CV…", expanded=False) as status:
        try:
            # Store raw bytes so we can re-open the original DOCX later for patching
            cv_bytes = uploaded_cv.getvalue()
            cv_text = extract_docx_text(io.BytesIO(cv_bytes))
            if not cv_text.strip():
                raise ValueError("No text found in the document.")
            status.update(label=f"CV read ({len(cv_text):,} chars)", state="complete")
        except Exception as exc:
            status.update(label="Could not read CV", state="error")
            st.error(str(exc))
            st.stop()

    with st.status("Analysing with Gemini AI…", expanded=False) as status:
        try:
            result = analyze_and_rewrite(cv_text, jd_text, gemini_api_key.strip())
            status.update(label="Analysis complete", state="complete")
        except json.JSONDecodeError as exc:
            status.update(label="Could not parse Gemini's response", state="error")
            st.error(f"JSON parse error: {exc}")
            st.stop()
        except Exception as exc:
            status.update(label="Gemini API error", state="error")
            st.error(str(exc))
            st.stop()

    st.session_state.analyzed = True
    st.session_state.result = result
    st.session_state.cv_text = cv_text
    st.session_state.cv_bytes = cv_bytes   # original DOCX bytes for patching
    st.session_state.jd_text = jd_text

# ─── Results ──────────────────────────────────────────────────────────────────

if not st.session_state.get("analyzed"):
    st.stop()

result   = st.session_state.result
cv_text  = st.session_state.cv_text
cv_bytes = st.session_state.cv_bytes
bullets  = result.get("bullets", [])
weak     = [b for b in bullets if b.get("is_weak")]
strong   = [b for b in bullets if not b.get("is_weak")]
score_before     = result.get("match_score_before", 0)
score_after      = result.get("match_score_after", 0)
ats_score_before = result.get("ats_score_before", 0)
ats_score_after  = result.get("ats_score_after", 0)
num_weak         = len(weak)

accepted_changes = [weak[i] for i in range(num_weak) if st.session_state.get(f"accept_{i}", True)]
num_accepted = len(accepted_changes)
total_weight    = sum(b.get("impact_score", 5) for b in weak)
accepted_weight = sum(weak[i].get("impact_score", 5) for i in range(num_weak)
                      if st.session_state.get(f"accept_{i}", True))
ratio          = (accepted_weight / total_weight) if total_weight else 1.0
live_score     = round(score_before + (score_after - score_before) * ratio)
live_ats_score = round(ats_score_before + (ats_score_after - ats_score_before) * ratio)

# ── Job title banner ──────────────────────────────────────────────────────────

job_title    = result.get("job_title", "")
company_name = result.get("company_name") or ""
role_label   = f"{job_title} at {company_name}" if company_name else job_title

st.divider()
if role_label:
    st.markdown(
        f"<p style='font-size:1.15rem;font-weight:600;margin-bottom:0.25rem'>"
        f"Analyzing your CV for: <span style='color:#1565c0'>{role_label}</span></p>",
        unsafe_allow_html=True,
    )

# ── Score metrics ─────────────────────────────────────────────────────────────

# Row 1 — four score metrics with narrow spacer columns as visual padding
r1c1, r1sp1, r1c2, r1mid, r1c3, r1sp2, r1c4 = st.columns([3, 0.4, 3, 0.6, 3, 0.4, 3])
r1c1.metric("JD Match — Before", f"{score_before}%")
r1sp1.markdown("<div style='text-align:center;font-size:1.6rem;padding-top:1.6rem;color:#9e9e9e'>→</div>", unsafe_allow_html=True)
r1c2.metric("JD Match — Live", f"{live_score}%",
            delta=f"+{live_score - score_before}%" if live_score >= score_before else f"{live_score - score_before}%")
r1mid.markdown("")
r1c3.metric("ATS Score — Before", f"{ats_score_before}%")
r1sp2.markdown("<div style='text-align:center;font-size:1.6rem;padding-top:1.6rem;color:#9e9e9e'>→</div>", unsafe_allow_html=True)
r1c4.metric("ATS Score — Live", f"{live_ats_score}%",
            delta=f"+{live_ats_score - ats_score_before}%" if live_ats_score >= ats_score_before else f"{live_ats_score - ats_score_before}%")

# Row 2 — secondary stats, centered under the score pairs
_, b1, _, b2, _ = st.columns([3, 3, 0.6, 3, 3])
b1.metric("Bullets Accepted", f"{num_accepted} / {num_weak}")
b2.metric("Strong Bullets", len(strong))

st.divider()

# ── Diff view with checkboxes ─────────────────────────────────────────────────

if not weak:
    st.success("Your CV is already well-aligned with this role — no rewrites needed.")
else:
    st.subheader(f"Suggested Rewrites ({num_weak})")
    st.caption("Tick the changes you want to keep. The score updates as you select.")

    _, hdr_l, hdr_r = st.columns([0.25, 2, 2])
    hdr_l.markdown('<div class="col-header">🔴 Original</div>', unsafe_allow_html=True)
    hdr_r.markdown('<div class="col-header">🟢 Rewritten</div>', unsafe_allow_html=True)

    for i, b in enumerate(weak):
        col_cb, col_l, col_r = st.columns([0.25, 2, 2])

        with col_cb:
            accepted = st.checkbox("Accept", key=f"accept_{i}", value=True,
                                   label_visibility="collapsed")
        with col_l:
            impact = b.get("impact_score", 5)
            impact_badge = f'<span class="impact-badge">Impact {impact}/10</span>'
            if accepted:
                st.markdown(
                    f'<div class="bullet-original">{impact_badge}{b["original"]}</div>'
                    f'<div class="bullet-reason">⚠ {b.get("reason", "")}</div>',
                    unsafe_allow_html=True,
                )
            else:
                st.markdown(
                    f'<div class="bullet-kept">{b["original"]}</div>'
                    f'<div class="bullet-reason">Keeping original</div>',
                    unsafe_allow_html=True,
                )
        with col_r:
            st.markdown(
                f'<div class="bullet-rewritten">{b["rewritten"]}</div>'
                if accepted else
                f'<div class="bullet-kept">{b["original"]}</div>',
                unsafe_allow_html=True,
            )
        st.write("")

# ── Strong bullets expander ───────────────────────────────────────────────────

if strong:
    with st.expander(f"✅ Strong bullets — keeping as-is ({len(strong)})"):
        for b in strong:
            st.markdown(f"**·** {b['original']}")
            if b.get("reason"):
                st.caption(b["reason"])

# ── Generate updated CV ───────────────────────────────────────────────────────

st.divider()

rejected_count = num_weak - num_accepted
st.markdown(
    f"**{num_accepted} change{'s' if num_accepted != 1 else ''} accepted"
    + (f", {rejected_count} rejected" if rejected_count else "")
    + "** — ready to rebuild your CV."
)

gen = st.button("📄 Generate My Updated CV", type="primary", use_container_width=True,
                disabled=(num_weak > 0 and num_accepted == 0))

if gen:
    for k in ("cv_generated", "word_bytes"):
        st.session_state.pop(k, None)

    with st.status("Applying changes to your original document…", expanded=False) as status:
        try:
            word_bytes = apply_changes_to_docx(io.BytesIO(cv_bytes), accepted_changes)
            status.update(label="Done — formatting preserved", state="complete")
        except Exception as exc:
            status.update(label="Failed to apply changes", state="error")
            st.error(str(exc))
            st.stop()

    st.session_state.cv_generated = True
    st.session_state.word_bytes = word_bytes

if st.session_state.get("cv_generated"):
    st.success(f"Your updated CV is ready — match score lifted from {score_before}% to {live_score}%.")
    st.download_button(
        label="⬇ Download Updated CV (.docx)",
        data=st.session_state.word_bytes,
        file_name="bespoke_cv.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )
